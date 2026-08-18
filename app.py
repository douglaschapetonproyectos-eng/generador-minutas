import streamlit as st
import pandas as pd
import math
import re
import io
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

try:
    import pypdf
    HAY_PYPDF = True
except:
    HAY_PYPDF = False

st.set_page_config(page_title="Generador de Minutas Topográficas", layout="wide", page_icon="📐")

def numero_a_letras(n):
    unidades = ["", "UN", "DOS", "TRES", "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE"]
    dieces = ["DIEZ", "ONCE", "DOCE", "TRECE", "CATORCE", "QUINCE", "DIECISÉIS", "DIECISIETE", "DIECIOCHO", "DIECINUEVE"]
    veintes = ["VEINTE", "VEINTIÚN", "VEINTIDÓS", "VEINTITRÉS", "VEINTICUATRO", "VEINTICINCO", "VEINTISÉIS", "VEINTISIETE", "VEINTIOCHO", "VEINTINUEVE"]
    decenas = ["", "DIEZ", "VEINTE", "TREINTA", "CUARENTA", "CINCUENTA", "SESENTA", "SETENTA", "OCHENTA", "NOVENTA"]
    centenas = ["", "CIENTO", "DOSCIENTOS", "TRESCIENTOS", "CUATROCIENTOS", "QUINIENTOS", "SEISCIENTOS", "SETECIENTOS", "OCHOCIENTOS", "NOVECIENTOS"]

    n_int = int(round(n))
    if n_int == 0:
        return "CERO"
    if n_int == 100:
        return "CIEN"

    def convertir_centena(num):
        if num == 0:
            return ""
        if num == 100:
            return "CIEN"
        c = num // 100
        d = (num % 100) // 10
        u = num % 10
        
        texto_c = centenas[c]
        resto = num % 100
        
        if resto == 0:
            return texto_c
        
        texto_resto = ""
        if 10 <= resto < 20:
            texto_resto = dieces[resto - 10]
        elif 20 <= resto < 30:
            texto_resto = veintes[resto - 20]
        else:
            if d > 0:
                texto_resto = decenas[d]
                if u > 0:
                    texto_resto += " Y " + unidades[u]
            else:
                texto_resto = unidades[u]
                
        if texto_c:
            return f"{texto_c} {texto_resto}".strip()
        return texto_resto.strip()

    millones = n_int // 1000000
    miles = (n_int % 1000000) // 1000
    resto = n_int % 1000

    partes = []
    if millones > 0:
        if millones == 1:
            partes.append("UN MILLÓN")
        else:
            partes.append(convertir_centena(millones) + " MILLONES")
    
    if miles > 0:
        if miles == 1:
            partes.append("MIL")
        else:
            partes.append(convertir_centena(miles) + " MIL")
            
    if resto > 0:
        partes.append(convertir_centena(resto))

    return " ".join(partes).strip()

def limpiar_numero(val):
    if pd.isna(val):
        return 0.0
    if isinstance(val, (int, float)):
        return float(val)
    s = str(val).strip()
    if not s or s.lower() in ['none', 'null', 'nan']:
        return 0.0
    if '.' in s and ',' in s:
        if s.rfind(',') > s.rfind('.'):
            s = s.replace('.', '').replace(',', '.')
        else:
            s = s.replace(',', '')
    elif ',' in s:
        s = s.replace(',', '.')
    s_clean = re.sub(r'[^\d.-]', '', s)
    try:
        return float(s_clean)
    except:
        return 0.0

def calcular_distancia(n1, e1, n2, e2):
    return math.sqrt((n2 - n1)**2 + (e2 - e1)**2)

def obtener_rumbo_cardinal(n1, e1, n2, e2):
    dn = n2 - n1
    de = e2 - e1
    if abs(dn) < 1e-7 and abs(de) < 1e-7:
        return "Mismo punto"
    rad = math.atan2(de, dn)
    grados = math.degrees(rad)
    if grados < 0:
        grados += 360
    
    if 337.5 <= grados or grados < 22.5:
        return "Norte"
    elif 22.5 <= grados < 67.5:
        return "Nor-Oriental"
    elif 67.5 <= grados < 112.5:
        return "Oriental"
    elif 112.5 <= grados < 157.5:
        return "Sur-Oriental"
    elif 157.5 <= grados < 202.5:
        return "Sur"
    elif 202.5 <= grados < 247.5:
        return "Sur-Occidental"
    elif 247.5 <= grados < 292.5:
        return "Occidental"
    else:
        return "Nor-Occidental"

def calcular_area_gauss(n, e):
    num_puntos = len(n)
    if num_puntos < 3:
        return 0.0
    area = 0.0
    for i in range(num_puntos):
        j = (i + 1) % num_puntos
        area += e[i] * n[j] - e[j] * n[i]
    return abs(area) / 2.0

def obtener_tramos_lindero(p_ini, p_fin, df, es_ultimo_cierre=False):
    df_c = df.reset_index(drop=True)
    n_p = len(df_c)
    pts_list = [str(x).strip() for x in df_c['Punto'].tolist()]
    
    if p_ini not in pts_list:
        return []
    
    idx_ini = pts_list.index(p_ini)
    seq = []
    
    if es_ultimo_cierre:
        curr = idx_ini
        while curr < n_p:
            seq.append(curr)
            curr += 1
        seq.append(0)  # Vuelve al punto inicial 1
    else:
        if p_fin not in pts_list:
            return []
        idx_fin = pts_list.index(p_fin)
        curr = idx_ini
        while True:
            seq.append(curr)
            if curr == idx_fin:
                break
            curr = (curr + 1) % n_p
            if curr == idx_ini:
                break
                
    tramos = []
    for k in range(len(seq) - 1):
        i1 = seq[k]
        i2 = seq[k+1]
        p1 = str(df_c.iloc[i1]['Punto']).strip()
        n1 = float(limpiar_numero(df_c.iloc[i1]['Norte']))
        e1 = float(limpiar_numero(df_c.iloc[i1]['Este']))
        p2 = str(df_c.iloc[i2]['Punto']).strip()
        n2 = float(limpiar_numero(df_c.iloc[i2]['Norte']))
        e2 = float(limpiar_numero(df_c.iloc[i2]['Este']))
        
        dist_calc = calcular_distancia(n1, e1, n2, e2)
        if 'Distancia' in df_c.columns and pd.notna(df_c.iloc[i1]['Distancia']) and limpiar_numero(df_c.iloc[i1]['Distancia']) > 0:
            dist = float(limpiar_numero(df_c.iloc[i1]['Distancia']))
        else:
            dist = dist_calc
            
        rumbo = obtener_rumbo_cardinal(n1, e1, n2, e2)
        tramos.append({
            "origen": p1, "destino": p2,
            "n1": n1, "e1": e1, "n2": n2, "e2": e2,
            "dist": dist, "rumbo": rumbo
        })
    return tramos

def redactar_lindero_item(costado_nombre, num_lindero, p_ini, p_fin, colindante, elemento, df, es_primer_del_costado=True, es_ultimo_cierre=False):
    tramos = obtener_tramos_lindero(p_ini, p_fin, df, es_ultimo_cierre=es_ultimo_cierre)
    if not tramos:
        return ""
    
    dist_total = sum(t['dist'] for t in tramos)
    p_orig = tramos[0]
    p_dest = tramos[-1]
    
    tipo_linea = "en línea semi-recta" if len(tramos) == 1 else "en línea quebrada"
    sentido = p_orig['rumbo']
    colind_txt = str(colindante).strip() if str(colindante).strip() else "predio vecino"
    elem_txt = str(elemento).strip() if str(elemento).strip() else "cerca de alambre al medio"
    
    if not elem_txt.endswith("al medio") and not elem_txt.startswith("quebrada") and not elem_txt.startswith("río") and not elem_txt.startswith("camino"):
        elem_txt_completo = f"{elem_txt} al medio."
    else:
        elem_txt_completo = f"{elem_txt}."
        
    prefijo = f"Por el {costado_nombre}: " if es_primer_del_costado else ""
    
    txt = f"{prefijo}lindero {num_lindero}: inicia en el mojón {p_orig['origen']} (N: {p_orig['n1']:.4f}, E: {p_orig['e1']:.4f}); "
    txt += f"en sentido {sentido} {tipo_linea}, "
    
    if len(tramos) == 1:
        if es_ultimo_cierre:
            txt += f"regresando hasta el mojón {p_dest['destino']} (N: {p_dest['n2']:.4f}, E: {p_dest['e2']:.4f}) punto de partida y encierra, "
        else:
            txt += f"hasta el mojón {p_dest['destino']} (N: {p_dest['n2']:.4f}, E: {p_dest['e2']:.4f}); "
    else:
        for t in tramos[:-1]:
            txt += f"hasta el mojón {t['destino']} (N: {t['n2']:.4f}, E: {t['e2']:.4f}); "
            
        if es_ultimo_cierre:
            txt += f"regresando hasta el mojón {p_dest['destino']} (N: {p_dest['n2']:.4f}, E: {p_dest['e2']:.4f}) punto de partida y encierra, "
        else:
            txt += f"hasta el mojón {p_dest['destino']} (N: {p_dest['n2']:.4f}, E: {p_dest['e2']:.4f}); "
            
    txt += f"colindando con el {colind_txt} {elem_txt_completo} teniendo como distancia total para este lindero de {dist_total:.2f} mts."
    return txt

st.title("📐 Generador Automatizado de Minutas Topográficas y Linderos")
st.markdown("Conforme al **Art. 2.2.2.2.17 del Decreto 148 de 2020** (Sistema MAGNA-SIRGAS Origen Único EPSG 9377).")

# 1. Datos Generales
with st.expander("📝 1. Identificación del Inmueble y Jurisdicción Registral", expanded=True):
    col1, col2, col3 = st.columns(3)
    with col1:
        nombre_predio = st.text_input("Nombre del Predio", value="SAN AGUSTIN")
        departamento = st.text_input("Departamento", value="Cundinamarca")
        municipio = st.text_input("Municipio", value="San Francisco")
    with col2:
        vereda = st.text_input("Vereda / Sector", value="El Arrayan")
        cedula_catastral = st.text_input("Número Predial / Cédula Catastral", value="256580000000000020232000000000")
        matricula_inmobiliaria = st.text_input("Matrícula Inmobiliaria", value="156-47224")
    with col3:
        circuito_registral = st.text_input("Circuito Registral / ORIP", value="Facatativá")
        profesional = st.text_input("Topógrafo / Perito Responsable", value="douglas chapeton")
        matricula_prof = st.text_input("Matrícula Profesional", value="01 19914 cpnt")

# 2. Coordenadas
st.markdown("### 📍 2. Coordenadas del Polígono (Excel / CSV o Tabla)")
archivo_coords = st.file_uploader("Cargar archivo Excel (.xlsx, .xls) o CSV con coordenadas", type=["xlsx", "xls", "csv"], key="uploader_coords")

df_cargado = pd.DataFrame(columns=["Punto", "Norte", "Este", "Distancia"])

if archivo_coords is not None:
    try:
        if archivo_coords.name.endswith('.csv'):
            df_raw = pd.read_csv(archivo_coords)
        else:
            df_raw = pd.read_excel(archivo_coords)
        
        cols = [str(c).strip() for c in df_raw.columns]
        df_raw.columns = cols
        
        def col_match(lista, ops):
            for c in lista:
                c_l = str(c).lower().strip()
                for o in ops:
                    if o in c_l:
                        return c
            return lista[0] if lista else None

        c_pto = col_match(cols, ['punto', 'pto', 'id', 'name', 'vertice', 'est', 'item', 'no', 'mojon'])
        c_nor = col_match(cols, ['norte', 'north', 'lat', 'y'])
        c_est = col_match(cols, ['este', 'east', 'lon', 'x'])
        c_dist = col_match(cols, ['distancia', 'dist', 'longitud', 'dist_m'])

        with st.expander("⚙️ Asignación de Columnas del Archivo"):
            sc1, sc2, sc3, sc4 = st.columns(4)
            sel_p = sc1.selectbox("Columna Punto / Mojón", cols, index=cols.index(c_pto) if c_pto in cols else 0)
            sel_n = sc2.selectbox("Columna Norte", cols, index=cols.index(c_nor) if c_nor in cols else 0)
            sel_e = sc3.selectbox("Columna Este", cols, index=cols.index(c_est) if c_est in cols else 0)
            opciones_dist = ["(Calcular automáticamente)"] + cols
            index_dist = (cols.index(c_dist) + 1) if (c_dist and c_dist in cols and c_dist not in [sel_p, sel_n, sel_e]) else 0
            sel_dist = sc4.selectbox("Columna Distancia (Opcional)", opciones_dist, index=index_dist)

        temp_df = pd.DataFrame()
        temp_df['Punto'] = df_raw[sel_p].astype(str).str.strip()
        temp_df['Norte'] = df_raw[sel_n].apply(limpiar_numero)
        temp_df['Este'] = df_raw[sel_e].apply(limpiar_numero)
        if sel_dist != "(Calcular automáticamente)":
            temp_df['Distancia'] = df_raw[sel_dist].apply(limpiar_numero)

        mask_valida = (
            (~temp_df['Punto'].str.upper().isin(['NONE', 'NAN', 'PUNTO', 'PTO', 'VERTICE', 'MOJON', 'ID', ''])) &
            (temp_df['Norte'] > 100) &
            (temp_df['Este'] > 100)
        )
        df_cargado = temp_df[mask_valida].reset_index(drop=True)
        st.success(f"✅ Se cargaron {len(df_cargado)} vértices válidos del archivo.")
    except Exception as e:
        st.error(f"Error al leer archivo: {e}")

st.markdown("#### Tabla de Coordenadas del Polígono:")
df_editor = st.data_editor(df_cargado, num_rows="dynamic", use_container_width=True)

df_limpio = df_editor.copy()
if not df_limpio.empty:
    df_limpio['Punto'] = df_limpio['Punto'].astype(str).str.strip()
    df_limpio['Norte'] = df_limpio['Norte'].apply(limpiar_numero)
    df_limpio['Este'] = df_limpio['Este'].apply(limpiar_numero)
    mask = (
        (~df_limpio['Punto'].str.upper().isin(['NONE', 'NAN', 'PUNTO', 'PTO', 'VERTICE', 'MOJON', 'ID', ''])) &
        (df_limpio['Norte'] > 100) &
        (df_limpio['Este'] > 100)
    )
    df_limpio = df_limpio[mask].reset_index(drop=True)

lista_puntos = [str(x).strip() for x in df_limpio['Punto'].tolist() if str(x).strip()]

# 3. Configuración de Linderos
st.markdown("### 🧭 3. Configuración de Colindancias por Costados Cardinales")

opciones_elementos = [
    "cerca de alambre",
    "cerca viva",
    "muro en ladrillo",
    "muro en piedra / mampostería",
    "vía pública",
    "camino veredal",
    "servidumbre de acceso",
    "quebrada aguas arriba",
    "quebrada aguas abajo",
    "río aguas arriba",
    "río aguas abajo",
    "mojones de concreto y línea imaginaria",
    "límite natural según levantamiento"
]

if len(lista_puntos) >= 3:
    config_linderos = []
    costados = ["Norte", "Oriente", "Sur", "Occidente"]
    
    for c_nombre in costados:
        st.markdown(f"#### 🌐 Costado {c_nombre}")
        num_colinds = st.number_input(f"Número de colindantes por el {c_nombre}:", min_value=1, max_value=10, value=1, key=f"num_{c_nombre}")
        
        for k in range(int(num_colinds)):
            sub_id = f"{c_nombre}_{k+1}"
            st.caption(f"Tramo / Colindante #{k+1} del Costado {c_nombre}")
            c1, c2, c3, c4 = st.columns(4)
            
            p_ini_def = lista_puntos[0] if k == 0 and c_nombre == "Norte" else lista_puntos[min(k, len(lista_puntos)-1)]
            p_fin_def = lista_puntos[min(k+1, len(lista_puntos)-1)]
            
            p_ini = c1.selectbox(f"Inicia en Mojón ({sub_id})", lista_puntos, index=lista_puntos.index(p_ini_def) if p_ini_def in lista_puntos else 0, key=f"ini_{sub_id}")
            
            es_ultimo = (c_nombre == "Occidente" and k == int(num_colinds) - 1)
            if es_ultimo:
                c2.info(f"Hasta: Mojón {lista_puntos[-1]} y cierra al Mojón {lista_puntos[0]}")
                p_fin = lista_puntos[-1]
            else:
                p_fin = c2.selectbox(f"Hasta Mojón ({sub_id})", lista_puntos, index=lista_puntos.index(p_fin_def) if p_fin_def in lista_puntos else min(1, len(lista_puntos)-1), key=f"fin_{sub_id}")
                
            colind = c3.text_input(f"Predio colindante ({sub_id})", value="00 00 0002 0233 000", key=f"col_{sub_id}")
            elem = c4.selectbox(f"Elemento delimitador ({sub_id})", opciones_elementos, index=0, key=f"elem_{sub_id}")
            
            config_linderos.append({
                "costado": c_nombre,
                "p_ini": p_ini,
                "p_fin": p_fin,
                "colindante": colind,
                "elemento": elem,
                "es_ultimo_cierre": es_ultimo
            })
else:
    st.info("👆 Carga o ingresa las coordenadas en la tabla superior para configurar los linderos.")

# 4. Anexo de Planos e Imágenes / PDF
st.markdown("### 🖼️ 4. Anexar Imágenes de Planos o Fotografías (PNG / JPG / PDF)")
archivos_planos = st.file_uploader("Adjuntar planos o fotos de linderos", type=["png", "jpg", "jpeg", "pdf"], accept_multiple_files=True, key="uploader_planos")

imagenes_para_word = []
if archivos_planos:
    cols_img = st.columns(min(len(archivos_planos), 3))
    for idx, f in enumerate(archivos_planos):
        col_act = cols_img[idx % 3]
        if f.type in ["image/png", "image/jpeg", "image/jpg"]:
            try:
                img = Image.open(f)
                col_act.image(img, caption=f.name, use_container_width=True)
                imagenes_para_word.append((f.name, f.getvalue()))
            except Exception as e:
                col_act.error(f"Error en imagen: {e}")
        elif f.type == "application/pdf":
            col_act.info(f"📄 Archivo PDF: **{f.name}** ({f.size / 1024:.1f} KB)")
            if HAY_PYPDF:
                try:
                    pdf_reader = pypdf.PdfReader(io.BytesIO(f.getvalue()))
                    col_act.caption(f"Páginas: {len(pdf_reader.pages)}")
                except:
                    pass

# 5. Generación de la Minuta
st.markdown("---")
if st.button("🚀 Generar Minuta Técnica Oficial y Documento Word", type="primary"):
    if len(df_limpio) < 3:
        st.error("Se requieren al menos 3 vértices con coordenadas válidas para calcular el polígono.")
    else:
        n_vals = df_limpio['Norte'].values
        e_vals = df_limpio['Este'].values
        n_puntos = len(df_limpio)

        tramos_completos = []
        perimetro_total = 0.0
        for i in range(n_puntos):
            sig = (i + 1) % n_puntos
            p1 = str(df_limpio.iloc[i]['Punto']).strip()
            n1 = float(n_vals[i])
            e1 = float(e_vals[i])
            p2 = str(df_limpio.iloc[sig]['Punto']).strip()
            n2 = float(n_vals[sig])
            e2 = float(e_vals[sig])
            
            dist_calc = calcular_distancia(n1, e1, n2, e2)
            if 'Distancia' in df_limpio.columns and pd.notna(df_limpio.iloc[i]['Distancia']) and limpiar_numero(df_limpio.iloc[i]['Distancia']) > 0:
                dist = float(limpiar_numero(df_limpio.iloc[i]['Distancia']))
            else:
                dist = dist_calc
                
            perimetro_total += dist
            rumbo_cardinal = obtener_rumbo_cardinal(n1, e1, n2, e2)
            
            tramos_completos.append({
                "Punto": p1,
                "Norte": n1,
                "Este": e1,
                "Destino": p2,
                "Distancia_m": dist,
                "Sentido_Rumbo": rumbo_cardinal
            })

        df_tabla_tramos = pd.DataFrame(tramos_completos)
        area_m2 = calcular_area_gauss(n_vals, e_vals)
        hectareas = int(area_m2 // 10000)
        metros_restantes = round(area_m2 % 10000, 2)
        area_en_letras = numero_a_letras(area_m2)

        st.success("✅ Minuta técnica redactada y cálculos generados con éxito.")

        m1, m2, m3 = st.columns(3)
        m1.metric("Área en Hectáreas", f"{hectareas} ha + {metros_restantes:,.0f} M²", f"{area_m2:,.2f} m² totales")
        m2.metric("Perímetro Total", f"{perimetro_total:,.2f} Metros")
        m3.metric("Área en Letras", f"{area_en_letras}")

        # Redacción de linderos
        cuerpo_linderos_lista = []
        num_lindero_contador = 1
        costado_anterior = None
        
        for item in config_linderos:
            es_primer = (item['costado'] != costado_anterior)
            costado_anterior = item['costado']
            
            txt_lindero = redactar_lindero_item(
                item['costado'],
                num_lindero_contador,
                item['p_ini'],
                item['p_fin'],
                item['colindante'],
                item['elemento'],
                df_limpio,
                es_primer_del_costado=es_primer,
                es_ultimo_cierre=item['es_ultimo_cierre']
            )
            if txt_lindero:
                cuerpo_linderos_lista.append(txt_lindero)
                num_lindero_contador += 1

        cuerpo_linderos = "\n\n".join(cuerpo_linderos_lista)

        encabezado_minuta = (
            f"MINUTA TECNICA DE LINDEROS\n\n"
            f"Corresponde a un inmueble identificado con el numero predial {cedula_catastral} ubicado en la zona Rural Vereda {vereda} del Municipio de {municipio}-{departamento}, denominado {nombre_predio.upper()} e inscrito en el circuito registral de {circuito_registral} bajo matricula inmobiliaria No {matricula_inmobiliaria}\n\n"
            f"La descripción técnica de los linderos del inmueble, en observancia del artículo 2.2.2.2.17 del decreto 148 de 2020, estableciendo la forma, orientación y extensión de los linderos en los siguientes términos:\n\n"
            f"“el bien inmueble identificado catastralmente con el numero predial {cedula_catastral} y folio de matrícula inmobiliaria {matricula_inmobiliaria}, denominado “{nombre_predio.upper()}”, presenta los siguientes linderos referidos al sistema de coordenadas proyectadas magna sirgas origen único nacional con epsg 9377:\n\n"
        )

        cierre_area = f"\n\nDe acuerdo con los anteriores linderos, el área del citado bien inmueble es de {area_m2:,.0f} metros cuadrados o {area_en_letras} metros cuadrados ({hectareas} ha + {metros_restantes:,.0f} M2).\n\n"
        
        bloque_firma = f"Profesional responsable:\n\n\n{profesional.lower()}\nmat prof: {matricula_prof.lower()}"

        texto_minuta_oficial = encabezado_minuta + cuerpo_linderos + cierre_area + bloque_firma

        st.subheader("📄 Minuta Técnica Oficial Generada (Calibri 12)")
        st.text_area("Texto oficial listo para copiar y pegar directamente en Word:", value=texto_minuta_oficial, height=450)

        st.subheader("📊 Cuadro Técnico de Coordenadas y Distancias")
        st.dataframe(df_tabla_tramos[["Punto", "Norte", "Este", "Destino", "Distancia_m", "Sentido_Rumbo"]], use_container_width=True)

        # Generar Documento Word (.docx) con estilo Calibri 12
        doc = Document()
        
        # Configuración de tipografía Calibri 12
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(12)
        
        titulo = doc.add_heading("MINUTA TECNICA DE LINDEROS", 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph(
            f"Corresponde a un inmueble identificado con el numero predial {cedula_catastral} ubicado en la zona Rural Vereda {vereda} del Municipio de {municipio}-{departamento}, denominado {nombre_predio.upper()} e inscrito en el circuito registral de {circuito_registral} bajo matricula inmobiliaria No {matricula_inmobiliaria}"
        )
        
        doc.add_paragraph(
            "La descripción técnica de los linderos del inmueble, en observancia del artículo 2.2.2.2.17 del decreto 148 de 2020, estableciendo la forma, orientación y extensión de los linderos en los siguientes términos:"
        )

        doc.add_paragraph(
            f"“el bien inmueble identificado catastralmente con el numero predial {cedula_catastral} y folio de matrícula inmobiliaria {matricula_inmobiliaria}, denominado “{nombre_predio.upper()}”, presenta los siguientes linderos referidos al sistema de coordenadas proyectadas magna sirgas origen único nacional con epsg 9377:"
        )

        for p_lind in cuerpo_linderos_lista:
            doc.add_paragraph(p_lind)
        
        p_cierre = doc.add_paragraph(f"\nDe acuerdo con los anteriores linderos, el área del citado bien inmueble es de {area_m2:,.0f} metros cuadrados o {area_en_letras} metros cuadrados ({hectareas} ha + {metros_restantes:,.0f} M2).")
        p_cierre.runs[0].bold = True

        p_firma_doc = doc.add_paragraph(f"\nProfesional responsable:\n\n\n{profesional.lower()}\nmat prof: {matricula_prof.lower()}")
        p_firma_doc.runs[0].bold = True

        doc.add_heading("CUADRO TÉCNICO DE COORDENADAS Y DISTANCIAS", level=1)
        tabla = doc.add_table(rows=1, cols=6)
        tabla.style = 'Table Grid'
        
        hdr_cells = tabla.rows[0].cells
        nombres_encabezados = ['Punto', 'Norte (m)', 'Este (m)', 'Destino', 'Distancia (m)', 'Sentido / Rumbo']
        for col_idx, col_name in enumerate(nombres_encabezados):
            hdr_cells[col_idx].text = col_name
        
        for t in tramos_completos:
            row_cells = tabla.add_row().cells
            fila_valores = [
                str(t['Punto']),
                f"{t['Norte']:,.4f}",
                f"{t['Este']:,.4f}",
                str(t['Destino']),
                f"{t['Distancia_m']:.2f}",
                str(t['Sentido_Rumbo'])
            ]
            for c_i, val_celda in enumerate(fila_valores):
                row_cells[c_i].text = val_celda

        # Anexos
        if imagenes_para_word:
            doc.add_page_break()
            doc.add_heading("ANEXOS CARTOGRÁFICOS Y FOTOGRÁFICOS", level=1)
            for nombre_img, bytes_img in imagenes_para_word:
                doc.add_paragraph(f"Plano / Fotografía: {nombre_img}")
                try:
                    img_stream = io.BytesIO(bytes_img)
                    doc.add_picture(img_stream, width=Inches(5.8))
                except Exception as ex_img:
                    doc.add_paragraph(f"[No se pudo incrustar imagen: {ex_img}]")
                doc.add_paragraph("")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        st.download_button(
            label="📥 Descargar Minuta Oficial en Word (.docx)",
            data=buffer,
            file_name=f"Minuta_{nombre_predio.replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
